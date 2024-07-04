

import docx

import os
import json

def check_folder_exists(folder_path):
  """
  Checks if a folder exists and returns an error message if it doesn't.

  Args:
      folder_path (str): The path to the folder to check.

  Returns:
      str: An error message if the folder doesn't exist, otherwise None.
  """
  if not os.path.isdir(folder_path):
    
    return False
  return True

def get_all_files(directory):
  

  """
  Gets all files in a directory recursively.

  Args:
    directory: The directory to search.

  Returns:
    A list of all files in the directory.
  """
  folderexit = check_folder_exists(directory)
  if not folderexit:
    print(f"Error: Folder '{directory}' does not exist.")
    return
  files = []
  for root, _, filenames in os.walk(directory):
    for filename in filenames:
      files.append(os.path.join(root, filename))
  return files


def getJsondatafromfile(filepath):
  with open(filepath, 'r') as f:
    # Code to read from the file using 'f' goes here
    # Example:
    data = json.load(f)
    return(data)
  

def migrateJsonToDocx(migrationFile):
    data = getJsondatafromfile(migrationFile)
    subject = data["subject"]
    examType= data["exam_type"]
    
    doc = docx.Document()
    doc.add_heading(f'{subject} {examType} Multiple Choice Questions', 0)
    for topic in data["topics"]:
      para1 = f'{topic["topic"]}'
      para2 = f'{topic["questions"]}'
      doc.add_paragraph(para1)
      doc.add_paragraph(para2)
      
      

    doc.save(f'output/{subject.lower()}.docx')


def main():
  migrationFiles = get_all_files("./jsonoutput")
  for file in migrationFiles:
    migrateJsonToDocx(file)
 


main()
     
